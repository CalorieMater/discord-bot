import os
import discord
import datetime
from docx import Document

print("🔁 起動開始")

# ==== 環境変数 ====
DISCORD_TOKEN = os.getenv("DISCORD_TOKEN")
print("🔐 DISCORD_TOKEN の読み込み:", "OK" if DISCORD_TOKEN else "❌ None")

# ==== Discord Intents ====
intents = discord.Intents.default()
intents.message_content = True
print("⚙️ Intents 設定完了")

client = discord.Client(intents=intents)
print("🧠 Discordクライアント作成完了")

# ==== チャンネルIDを設定（数値にするのを忘れずに！） ====
TARGET_CHANNEL_ID = int(os.getenv("TARGET_CHANNEL_ID", "0"))  # 例: 123456789012345678

# ==== 月ごとのファイル名作成 ====
def get_monthly_filename(prefix, extension, year=None, month=None):
    if year is None or month is None:
        now = datetime.datetime.now()
        year, month = now.year, now.month
    return f"{prefix}_{year}_{month:02d}.{extension}"

# ==== メッセージ保存 ====
def save_message(user, content):
    now = datetime.datetime.now()
    timestamp = now.strftime("%Y-%m-%d %H:%M:%S")
    txt_filename = get_monthly_filename("summaries", "txt")
    with open(txt_filename, "a", encoding="utf-8") as f:
        f.write(f"[{timestamp}] {user}: {content}\n")

# ==== .txt → .docx 変換 ====
def convert_txt_to_docx(year=None, month=None):
    txt_filename = get_monthly_filename("summaries", "txt", year, month)
    docx_filename = get_monthly_filename("summary", "docx", year, month)

    if not os.path.exists(txt_filename):
        return None

    document = Document()
    document.add_heading("月次メッセージログ", level=1)

    with open(txt_filename, "r", encoding="utf-8") as f:
        for line in f:
            document.add_paragraph(line.strip())

    document.save(docx_filename)
    return docx_filename

# ==== Bot 起動時 ====
@client.event
async def on_ready():
    print(f"✅ Bot 起動完了: {client.user}")

# ==== メッセージ受信時 ====
@client.event
async def on_message(message):
    if message.author.bot:
        return

    if message.channel.id != TARGET_CHANNEL_ID:
        return

    content = message.content.strip()

    # ✅ !send-summary [YYYY-MM]
    if content.startswith("!send-summary"):
        parts = content.split()
        if len(parts) == 2:
            try:
                year, month = map(int, parts[1].split("-"))
                docx_filename = get_monthly_filename("summary", "docx", year, month)
            except ValueError:
                await message.channel.send("⚠️ コマンド形式が正しくありません！例: `!send-summary 2025-03`")
                return
        else:
            now = datetime.datetime.now()
            docx_filename = get_monthly_filename("summary", "docx", now.year, now.month)

        if os.path.exists(docx_filename):
            await message.channel.send(f"📄 {docx_filename} を送信します：", file=discord.File(docx_filename))
        else:
            await message.channel.send(f"⚠️ ファイルが見つかりません: `{docx_filename}`")

    else:
        save_message(message.author.display_name, content)
        print(f"📝 保存: {message.author.display_name}: {content}")

# ==== Bot 起動 ====
try:
    client.run(DISCORD_TOKEN)
except Exception as e:
    print("❌ 起動エラー:", e)